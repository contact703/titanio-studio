#!/usr/bin/env python3
"""
preencher-edital.py — Motor de preenchimento inteligente de editais para a Titanio Studio.
Victor Capital | Titanio Studio

Uso:
    python preencher-edital.py --analise analise.json
    python preencher-edital.py --analise analise.json --formato docx
    python preencher-edital.py --analise analise.json --formato json --saida preenchimento.json
"""

import sys
import os
import re
import json
import argparse
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from perfil_titanio import PERFIL_TITANIO, carregar_perfil_titanio, selecionar_produto_relevante

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
MODELO_IA = "claude-opus-4-5"
PASTA_SAIDAS = Path(__file__).parent.parent / "saidas"
PASTA_SAIDAS.mkdir(exist_ok=True)

# ─── REGRAS DE TOM POR TIPO DE PROGRAMA ───────────────────────────────────────

REGRAS_TOM = {
    "FINEP": {
        "tom": "técnico-científico",
        "estilo": "formal, metodológico, com evidências quantitativas",
        "palavras_chave": ["inovação tecnológica", "P&D", "pesquisa aplicada", "metodologia", "resultados mensuráveis"],
        "foco": "capacidade técnica, metodologia de desenvolvimento, propriedade intelectual",
        "evitar": "linguagem de marketing, promessas sem evidência, jargão de startup",
        "instrucoes": (
            "Use linguagem técnica e científica. Cite metodologias, métricas de desenvolvimento, "
            "propriedade intelectual e impacto tecnológico. Estruture como proposta de P&D."
        ),
    },
    "CNPq": {
        "tom": "científico-acadêmico",
        "estilo": "rigoroso, metodológico, referências a literatura",
        "palavras_chave": ["pesquisa", "metodologia", "hipótese", "resultados", "publicações"],
        "foco": "contribuição científica, metodologia rigorosa, qualificação da equipe",
        "evitar": "foco excessivo em negócio/receita",
        "instrucoes": (
            "Foco na contribuição científica e metodologia. Destaque qualificação do fundador "
            "(Sorbonne), produções como validação empírica, alcance como evidência de impacto."
        ),
    },
    "BNDES": {
        "tom": "econômico-empresarial",
        "estilo": "formal, focado em retorno econômico e escalabilidade",
        "palavras_chave": ["geração de empregos", "receita", "escalabilidade", "retorno", "desenvolvimento"],
        "foco": "viabilidade econômica, geração de empregos, retorno do investimento",
        "evitar": "impacto social sem dados econômicos",
        "instrucoes": (
            "Destaque faturamento, capacidade de escala, geração de empregos, "
            "retorno econômico para o Brasil e potencial de exportação."
        ),
    },
    "BID": {
        "tom": "impacto-social",
        "estilo": "metrics-driven, Theory of Change, IRIS+",
        "palavras_chave": ["impacto social", "ODS", "beneficiários", "métricas IRIS+", "Theory of Change"],
        "foco": "impacto mensurável, beneficiários diretos/indiretos, ODS alinhados",
        "evitar": "foco apenas em negócio sem impacto social",
        "instrucoes": (
            "Construa uma Theory of Change clara. Use métricas IRIS+. "
            "Destaque VoxDescriber: 6,5M brasileiros com deficiência visual como beneficiários potenciais. "
            "Alinhe aos ODS relevantes."
        ),
    },
    "YC": {
        "tom": "startup-direto",
        "estilo": "conciso, confiante, focado em tração e mercado",
        "palavras_chave": ["tração", "crescimento", "mercado", "produto-mercado", "ARR"],
        "foco": "crescimento, tração, tamanho de mercado, equipe excepcional",
        "evitar": "linguagem corporativa, burocracia, detalhes técnicos excessivos",
        "instrucoes": (
            "Seja direto e confiante. Lidere com tração (470M espectadores, 40 países). "
            "Mostre tamanho de mercado, vantagem competitiva e por que agora é o momento."
        ),
    },
    "Aceleradora": {
        "tom": "pitch-startup",
        "estilo": "energético, focado em crescimento, produto-mercado fit",
        "palavras_chave": ["tração", "escala", "produto", "mercado", "crescimento"],
        "foco": "produto, mercado, equipe, tração, potencial de crescimento",
        "evitar": "história muito longa, detalhes técnicos sem impacto no negócio",
        "instrucoes": (
            "Pitch claro: Problema → Solução → Mercado → Tração → Equipe → Ask. "
            "Destaque 20 anos de tração no audiovisual + pivot para IA como diferencial único."
        ),
    },
    "ANCINE": {
        "tom": "audiovisual-cultural",
        "estilo": "técnico-cultural, foco em conteúdo e distribuição",
        "palavras_chave": ["produção audiovisual", "distribuição", "conteúdo", "obra", "janelas"],
        "foco": "qualidade da produção, distribuição, público, impacto cultural",
        "evitar": "foco excessivo em IA sem conexão com audiovisual",
        "instrucoes": (
            "Posicione como produtora audiovisual premium com 20 anos. "
            "Destaque o portfólio, distribuição internacional e alcance. "
            "IA é diferencial para produção, não o produto principal."
        ),
    },
    "Padrão": {
        "tom": "profissional-adaptado",
        "estilo": "claro, profissional, adaptado ao contexto",
        "palavras_chave": ["inovação", "impacto", "experiência", "resultados"],
        "foco": "histórico, capacidade de execução, relevância para o edital",
        "evitar": "generalidades sem evidência",
        "instrucoes": (
            "Use linguagem profissional e adapte ao contexto do edital. "
            "Sempre evidencie com dados: 20 anos, 40 produções, 40 países, 470M espectadores."
        ),
    },
}


# ─── MONTAGEM DO PROMPT DE PREENCHIMENTO ──────────────────────────────────────

def _montar_prompt_preenchimento(analise: dict, perfil: dict) -> str:
    tipo_programa = analise.get("tipo_programa", "Padrão")
    regras = REGRAS_TOM.get(tipo_programa, REGRAS_TOM["Padrão"])

    produto_key = analise.get("produto_recomendado", "cinema")
    produto = perfil["produtos"].get(produto_key, perfil["produtos"]["cinema"])

    angulo_key = analise.get("angulo_recomendado", "empresa_tecnologia")
    angulo = perfil["angulos_posicionamento"].get(angulo_key, "")

    criterios = analise.get("criterios_avaliacao", [])
    criterios_str = "\n".join([
        f"- {c.get('criterio', '')}: {c.get('descricao', '')} (peso: {c.get('peso', '?')})"
        for c in criterios
    ])

    docs_necessarios = analise.get("documentos_obrigatorios", [])

    return f"""Você é Victor Capital, especialista em captação de recursos da Titanio Studio.

## EDITAL A PREENCHER
- Título: {analise.get('titulo', '')}
- Organização: {analise.get('organizacao', '')}
- Valor máximo: R$ {analise.get('valor_max', 0):,.2f}
- Prazo: {analise.get('prazo_inscricao', '')}
- Tipo de programa: {tipo_programa}

## TOM RECOMENDADO: {regras['tom']}
{regras['instrucoes']}

## PRODUTO EM DESTAQUE PARA ESTE EDITAL: {produto.get('nome', '')}
{produto.get('descricao', '')}

## ÂNGULO DE POSICIONAMENTO:
{angulo}

## PONTOS FORTES A DESTACAR:
{chr(10).join('- ' + p for p in analise.get('pontos_fortes', []))}

## CRITÉRIOS DE AVALIAÇÃO:
{criterios_str if criterios_str else 'Verificar no edital'}

## PERFIL COMPLETO DA EMPRESA
- Nome: {perfil['nome']}
- CNPJ: {perfil.get('cnpj', 'a confirmar')}
- Fundação: {perfil['fundacao']}
- Sede: {perfil['sede']}
- Escritório internacional: {perfil['escritorio_internacional']}
- Fundador: {perfil['fundador']} — {perfil['formacao_fundador']}
- Anos de atividade: {perfil['historico']['anos_atividade']}
- Produções: {perfil['historico']['producoes']}+
- Países: {perfil['historico']['paises_distribuicao']}+
- Espectadores: {perfil['historico']['espectadores_formatado']}
- Pivot IA 2024: {perfil['pivot_ia_2024']}

## DOCUMENTOS OBRIGATÓRIOS (que precisarão ser preparados):
{chr(10).join('- ' + d for d in docs_necessarios)}

---

## SUA TAREFA
Gere o preenchimento COMPLETO e PERSONALIZADO desta inscrição. Retorne JSON com esta estrutura:

{{
  "titulo_edital": "{analise.get('titulo', '')}",
  "campos_preenchidos": [
    {{
      "campo": "nome do campo",
      "resposta": "texto completo preenchido",
      "caracteres": 0,
      "notas": "observações internas"
    }}
  ],
  "descricao_projeto": "texto completo para campo de descrição do projeto (500-1000 palavras)",
  "justificativa": "texto completo de justificativa (300-600 palavras)",
  "objetivos": {{
    "geral": "objetivo geral do projeto",
    "especificos": ["objetivo específico 1", "objetivo específico 2", "objetivo específico 3"]
  }},
  "metodologia": "descrição da metodologia de execução (300-500 palavras)",
  "cronograma": [
    {{"mes": 1, "atividade": "nome", "descricao": "o que será feito"}}
  ],
  "orcamento": {{
    "total": 0.0,
    "itens": [
      {{"categoria": "nome", "valor": 0.0, "justificativa": "por que é necessário"}}
    ]
  }},
  "resultados_esperados": ["resultado 1", "resultado 2"],
  "impacto_social": "descrição do impacto social (se aplicável)",
  "equipe": [
    {{"nome": "nome", "cargo": "função no projeto", "qualificacao": "experiência relevante"}}
  ],
  "documentos_checklist": [
    {{"documento": "nome do doc", "status": "preparar", "observacao": "onde obter/como preparar"}}
  ],
  "resumo_executivo": "versão de 3 parágrafos para revisar rapidamente",
  "alertas": ["pontos que precisam de revisão humana antes de submeter"]
}}

IMPORTANTE:
- Cada campo deve ter texto REAL e ESPECÍFICO para este edital
- Use os dados reais da Titanio (não invente números)
- Adapte o tom para: {regras['tom']}
- Destaque o produto: {produto.get('nome', '')}
- Orçamento total deve ser ≤ R$ {analise.get('valor_max', 0):,.2f}
- Inclua alertas para qualquer campo que precise de dado real (CNPJ, faturamento, etc.)

Retorne APENAS o JSON, sem texto adicional."""


# ─── PREENCHIMENTO COM IA ──────────────────────────────────────────────────────

def preencher_com_ia(analise: dict, perfil: dict) -> dict:
    """Usa Claude para gerar preenchimento completo do edital."""
    if not HAS_ANTHROPIC or not ANTHROPIC_API_KEY:
        print("⚠️  IA não disponível. Gerando preenchimento básico.")
        return _preenchimento_basico(analise, perfil)

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    prompt = _montar_prompt_preenchimento(analise, perfil)

    print(f"✍️  Gerando preenchimento para: {analise.get('titulo', 'Edital')}")
    print(f"🎨 Tom: {analise.get('tom_recomendado', 'profissional')}")
    print(f"🎯 Produto em foco: {analise.get('produto_recomendado', 'cinema')}")

    mensagem = client.messages.create(
        model=MODELO_IA,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )

    resposta = mensagem.content[0].text.strip()

    try:
        resultado = json.loads(resposta)
    except json.JSONDecodeError:
        match = re.search(r'\{.*\}', resposta, re.DOTALL)
        if match:
            resultado = json.loads(match.group())
        else:
            raise ValueError(f"IA retornou resposta inválida")

    return resultado


def _preenchimento_basico(analise: dict, perfil: dict) -> dict:
    """Preenchimento básico sem IA — estrutura com dados do perfil."""
    produto_key = analise.get("produto_recomendado", "cinema")
    produto = perfil["produtos"].get(produto_key, perfil["produtos"]["cinema"])

    historico = perfil["historico"]
    fundador = perfil["fundador"]

    return {
        "titulo_edital": analise.get("titulo", ""),
        "campos_preenchidos": [
            {
                "campo": "Nome da empresa",
                "resposta": perfil["nome"],
                "caracteres": len(perfil["nome"]),
                "notas": "",
            },
            {
                "campo": "Descrição da empresa",
                "resposta": (
                    f"{perfil['nome']} é uma empresa de produção audiovisual e tecnologia "
                    f"fundada em {perfil['fundacao']} por {fundador}. "
                    f"Com sede em {perfil['sede']} e escritório em {perfil['escritorio_internacional']}, "
                    f"a empresa acumula {historico['anos_atividade']} anos de experiência, "
                    f"{historico['producoes']}+ produções distribuídas em {historico['paises_distribuicao']}+ países, "
                    f"alcançando {historico['espectadores_formatado']} de espectadores."
                ),
                "caracteres": 500,
                "notas": "Adaptar ao limite de caracteres do campo",
            },
        ],
        "descricao_projeto": (
            f"O presente projeto propõe o desenvolvimento e expansão de {produto.get('nome', '')}, "
            f"{produto.get('descricao', '')} "
            f"A Titanio Studio, com {historico['anos_atividade']} anos de experiência e presença em "
            f"{historico['paises_distribuicao']} países, possui a infraestrutura e expertise necessárias "
            f"para executar com excelência este projeto de alto impacto."
        ),
        "justificativa": (
            f"A Titanio Studio é uma das poucas empresas brasileiras com capacidade de unir "
            f"tradição audiovisual ({historico['producoes']}+ produções, {historico['espectadores_formatado']} espectadores) "
            f"com inovação tecnológica de ponta. O projeto se alinha diretamente aos objetivos do edital "
            f"e representa uma oportunidade única de impacto nacional e internacional."
        ),
        "objetivos": {
            "geral": f"Desenvolver e escalar {produto.get('nome', '')} para maximizar impacto e alcance.",
            "especificos": [
                f"Expandir a base de usuários de {produto.get('nome', '')}",
                "Estabelecer parcerias estratégicas nacionais e internacionais",
                "Gerar impacto mensurável para os beneficiários do projeto",
            ],
        },
        "metodologia": (
            "A execução seguirá metodologia ágil em sprints quinzenais, com acompanhamento "
            "de KPIs mensuráveis e entregas parciais verificáveis. A equipe multidisciplinar "
            "da Titanio Studio, com expertise em audiovisual e tecnologia, garante capacidade "
            "de execução dentro do prazo e orçamento estabelecidos."
        ),
        "cronograma": [
            {"mes": 1, "atividade": "Planejamento e Setup", "descricao": "Definição de equipe, infraestrutura e marcos do projeto"},
            {"mes": 2, "atividade": "Desenvolvimento", "descricao": "Execução das atividades principais"},
            {"mes": 3, "atividade": "Testes e Validação", "descricao": "Validação com usuários e ajustes"},
            {"mes": 4, "atividade": "Lançamento", "descricao": "Lançamento e estratégia de distribuição"},
            {"mes": 5, "atividade": "Acompanhamento", "descricao": "Monitoramento de métricas e ajustes"},
            {"mes": 6, "atividade": "Relatório Final", "descricao": "Prestação de contas e relatório de impacto"},
        ],
        "orcamento": {
            "total": analise.get("valor_max", 0) * 0.85,
            "itens": [
                {"categoria": "Recursos Humanos", "valor": analise.get("valor_max", 0) * 0.50, "justificativa": "Equipe técnica e criativa"},
                {"categoria": "Tecnologia e Infraestrutura", "valor": analise.get("valor_max", 0) * 0.20, "justificativa": "Servidores, APIs, ferramentas"},
                {"categoria": "Pesquisa e Desenvolvimento", "valor": analise.get("valor_max", 0) * 0.10, "justificativa": "P&D e validação"},
                {"categoria": "Marketing e Distribuição", "valor": analise.get("valor_max", 0) * 0.05, "justificativa": "Divulgação e alcance"},
            ],
        },
        "resultados_esperados": [
            f"Expansão do alcance de {produto.get('nome', '')} em 300%",
            "Impacto direto em comunidades beneficiárias",
            "Retorno econômico e social mensurável",
        ],
        "impacto_social": produto.get("descricao", ""),
        "equipe": [
            {
                "nome": "Tiago Arakilian",
                "cargo": "Diretor Geral / Fundador",
                "qualificacao": f"Fundador da Titanio Studio, formado pela Sorbonne (Paris), {historico['anos_atividade']} anos de experiência",
            },
        ],
        "documentos_checklist": [
            {"documento": doc, "status": "preparar", "observacao": "Verificar requisitos específicos no edital"}
            for doc in analise.get("documentos_obrigatorios", [])
        ],
        "resumo_executivo": (
            f"A Titanio Studio solicita apoio para o projeto {analise.get('titulo', '')}.\n\n"
            f"Com {historico['anos_atividade']} anos de experiência, {historico['producoes']}+ produções em "
            f"{historico['paises_distribuicao']} países e {historico['espectadores_formatado']} de espectadores, "
            f"a empresa possui histórico comprovado para executar projetos de alto impacto.\n\n"
            f"O projeto proposto ({produto.get('nome', '')}) tem potencial de transformação "
            f"real para os beneficiários e alinhamento direto com os objetivos do edital."
        ),
        "alertas": [
            "⚠️  Revisar todos os valores de orçamento com a equipe financeira",
            "⚠️  Confirmar CNPJ e dados cadastrais antes de submeter",
            "⚠️  Verificar se todos os documentos obrigatórios estão prontos",
            "⚠️  Adaptar textos ao limite de caracteres de cada campo do sistema",
            "⚠️  Revisão final obrigatória antes do envio",
        ],
    }


# ─── FUNÇÃO PRINCIPAL ──────────────────────────────────────────────────────────

def preencher_edital(analise: dict, perfil_empresa: dict = None) -> dict:
    """
    Gera preenchimento completo e personalizado para o edital analisado.

    Args:
        analise: dict retornado por analisar_edital()
        perfil_empresa: dict do perfil da empresa (usa Titanio por padrão)

    Returns:
        dict com preenchimento completo
    """
    if perfil_empresa is None:
        perfil_empresa = carregar_perfil_titanio()

    preenchimento = preencher_com_ia(analise, perfil_empresa)

    # Adiciona metadados
    preenchimento["_meta"] = {
        "data_preenchimento": datetime.now().isoformat(),
        "edital_titulo": analise.get("titulo", ""),
        "score_fit": analise.get("fit_titanio_score", 0),
        "produto_usado": analise.get("produto_recomendado", ""),
        "tom_usado": analise.get("tom_recomendado", ""),
        "versao": "1.0",
        "status": "RASCUNHO — REVISÃO HUMANA OBRIGATÓRIA ANTES DE SUBMETER",
    }

    return preenchimento


def salvar_preenchimento(preenchimento: dict, titulo_edital: str = None, formato: str = "json") -> str:
    """Salva o preenchimento em arquivo JSON e/ou Word."""
    titulo_slug = re.sub(r'[^\w\s-]', '', titulo_edital or preenchimento.get("titulo_edital", "edital"))
    titulo_slug = re.sub(r'[-\s]+', '-', titulo_slug).strip('-').lower()[:50]
    data = datetime.now().strftime("%Y%m%d_%H%M")

    caminho_json = str(PASTA_SAIDAS / f"preenchimento_{titulo_slug}_{data}.json")
    with open(caminho_json, "w", encoding="utf-8") as f:
        json.dump(preenchimento, f, ensure_ascii=False, indent=2)
    print(f"💾 Preenchimento JSON salvo: {caminho_json}")

    if formato == "docx":
        caminho_docx = str(PASTA_SAIDAS / f"preenchimento_{titulo_slug}_{data}.docx")
        _gerar_docx(preenchimento, caminho_docx)
        return caminho_docx

    return caminho_json


def _gerar_docx(preenchimento: dict, caminho: str):
    """Gera documento Word com o preenchimento."""
    if not HAS_DOCX:
        print("⚠️  python-docx não instalado. pip install python-docx")
        return

    doc = Document()

    # Cabeçalho
    titulo = doc.add_heading("INSCRIÇÃO — " + preenchimento.get("titulo_edital", ""), 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = preenchimento.get("_meta", {})
    doc.add_paragraph(f"Data: {meta.get('data_preenchimento', '')[:10]}")
    doc.add_paragraph(f"Score de fit: {meta.get('score_fit', 0)}/100")
    doc.add_paragraph(f"STATUS: {meta.get('status', '')}")
    doc.add_paragraph()

    # Alertas em destaque
    alertas = preenchimento.get("alertas", [])
    if alertas:
        doc.add_heading("⚠️ ALERTAS ANTES DE SUBMETER", level=1)
        for alerta in alertas:
            p = doc.add_paragraph(alerta)
            p.runs[0].bold = True

    # Resumo executivo
    doc.add_heading("Resumo Executivo", level=1)
    doc.add_paragraph(preenchimento.get("resumo_executivo", ""))

    # Campos preenchidos
    doc.add_heading("Campos do Formulário", level=1)
    for campo in preenchimento.get("campos_preenchidos", []):
        doc.add_heading(campo.get("campo", ""), level=2)
        doc.add_paragraph(campo.get("resposta", ""))
        if campo.get("notas"):
            p = doc.add_paragraph(f"[Nota interna: {campo['notas']}]")
            p.runs[0].italic = True

    # Descrição do projeto
    doc.add_heading("Descrição do Projeto", level=1)
    doc.add_paragraph(preenchimento.get("descricao_projeto", ""))

    # Justificativa
    doc.add_heading("Justificativa", level=1)
    doc.add_paragraph(preenchimento.get("justificativa", ""))

    # Objetivos
    doc.add_heading("Objetivos", level=1)
    objetivos = preenchimento.get("objetivos", {})
    doc.add_paragraph(f"Geral: {objetivos.get('geral', '')}")
    doc.add_paragraph("Específicos:")
    for obj in objetivos.get("especificos", []):
        doc.add_paragraph(f"• {obj}", style="List Bullet")

    # Metodologia
    doc.add_heading("Metodologia", level=1)
    doc.add_paragraph(preenchimento.get("metodologia", ""))

    # Orçamento
    doc.add_heading("Orçamento", level=1)
    orcamento = preenchimento.get("orcamento", {})
    doc.add_paragraph(f"Total: R$ {orcamento.get('total', 0):,.2f}")
    for item in orcamento.get("itens", []):
        doc.add_paragraph(
            f"• {item['categoria']}: R$ {item['valor']:,.2f} — {item['justificativa']}",
            style="List Bullet"
        )

    # Documentos checklist
    doc.add_heading("Documentos Necessários", level=1)
    for doc_item in preenchimento.get("documentos_checklist", []):
        doc.add_paragraph(
            f"[ ] {doc_item['documento']} — {doc_item.get('observacao', '')}",
            style="List Bullet"
        )

    doc.save(caminho)
    print(f"📄 Documento Word salvo: {caminho}")


def imprimir_resumo_preenchimento(preenchimento: dict):
    """Imprime resumo do preenchimento gerado."""
    meta = preenchimento.get("_meta", {})
    campos = preenchimento.get("campos_preenchidos", [])
    alertas = preenchimento.get("alertas", [])

    print("\n" + "═" * 60)
    print("✍️  PREENCHIMENTO GERADO — Victor Capital")
    print("═" * 60)
    print(f"📋 Edital: {preenchimento.get('titulo_edital', 'N/A')}")
    print(f"🎯 Produto em foco: {meta.get('produto_usado', 'N/A')}")
    print(f"💬 Tom utilizado: {meta.get('tom_usado', 'N/A')}")
    print(f"📝 Campos preenchidos: {len(campos)}")
    print(f"\n📌 Resumo Executivo:")
    print(preenchimento.get("resumo_executivo", "N/A")[:300] + "...")

    if alertas:
        print(f"\n⚠️  ALERTAS ({len(alertas)}):")
        for a in alertas:
            print(f"   {a}")

    docs = preenchimento.get("documentos_checklist", [])
    if docs:
        print(f"\n📎 Documentos a preparar ({len(docs)}):")
        for d in docs:
            print(f"   [ ] {d['documento']}")

    print(f"\n🚨 STATUS: {meta.get('status', 'REVISAR ANTES DE SUBMETER')}")
    print("═" * 60)


# ─── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Victor Capital — Motor de preenchimento de editais para a Titanio Studio"
    )
    parser.add_argument("--analise", required=True, help="Caminho para arquivo JSON de análise")
    parser.add_argument("--formato", choices=["json", "docx"], default="json", help="Formato de saída")
    parser.add_argument("--saida", help="Caminho para salvar preenchimento")

    args = parser.parse_args()

    # Carrega análise
    with open(args.analise, "r", encoding="utf-8") as f:
        analise = json.load(f)

    perfil = carregar_perfil_titanio()
    preenchimento = preencher_edital(analise, perfil)
    imprimir_resumo_preenchimento(preenchimento)
    salvar_preenchimento(preenchimento, analise.get("titulo", ""), args.formato)


if __name__ == "__main__":
    main()
